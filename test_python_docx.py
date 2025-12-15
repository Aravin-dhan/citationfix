from docx import Document
import re
import os

def test_generation():
    text = "This is a test paragraph with a citation{{fn: Smith v. Jones, 123 F.3d 456}}. And another one{{fn: Doe v. Roe, 789 U.S. 101}}."
    
    print("Generating docx from text:", text)
    
    doc = Document()
    p = doc.add_paragraph()
    
    parts = re.split(r'(\{\{fn:.*?\}\})', text)
    
    for part in parts:
        if part.startswith('{{fn:') and part.endswith('}}'):
            citation = part[5:-2].strip()
            if citation:
                print(f"Found citation: {citation}")
                run = p.add_run()
                footnote = run.add_footnote()
                footnote.paragraphs[0].add_run(citation)
        else:
            if part:
                p.add_run(part)
                
    output_file = "test_output.docx"
    doc.save(output_file)
    print(f"Saved to {output_file}")
    
    if os.path.exists(output_file):
        print("SUCCESS: File generated.")
    else:
        print("FAILURE: File not generated.")

if __name__ == "__main__":
    test_generation()
