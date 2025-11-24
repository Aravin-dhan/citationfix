from http.server import BaseHTTPRequestHandler
import json
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, ns
import re

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_num_pages(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "NUMPAGES"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            # Read request body
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            data = json.loads(body)
            text = data.get('text', '')
            apply_formatting = data.get('formatting', False)
            convert_citations = data.get('convert_citations', True)

            # Validate input
            if not text:
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': 'No text provided'}).encode())
                return

            # Word count validation
            word_count = len(text.split())
            if word_count > 10000:
                self.send_response(400)
                self.send_header('Content-type', 'application/json')
                self.end_headers()
                self.wfile.write(json.dumps({'error': f'Text exceeds 10,000 word limit (current: {word_count})'}).encode())
                return

            # Process text and generate docx
            doc = Document()
            
            # Set default style to Times New Roman 12 if formatting enabled
            style = doc.styles['Normal']
            if apply_formatting:
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
            
            # Configure Footnote Text style if formatting is requested
            if apply_formatting:
                try:
                    footnote_style = doc.styles['Footnote Text']
                    footnote_style.font.name = 'Times New Roman'
                    footnote_style.font.size = Pt(10)
                    footnote_style.paragraph_format.line_spacing = 1.0
                except KeyError:
                    pass

            # Split text by newlines to preserve paragraphs
            paragraphs = text.split('\n')
            
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                    
                p = doc.add_paragraph()
                
                # Apply formatting to paragraph
                if apply_formatting:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5
                    p.paragraph_format.space_after = Pt(0)
                
                if convert_citations:
                    # Regex to find {{fn: ...}} markers
                    parts = re.split(r'(\{\{fn:.*?\}\})', para_text)
                    
                    for part in parts:
                        if part.startswith('{{fn:') and part.endswith('}}'):
                            # Extract citation text
                            citation = part[5:-2].strip()
                            if citation:
                                # Add footnote
                                footnote = p.add_footnote(citation)
                                
                                # Format the footnote text if needed
                                if apply_formatting:
                                    if footnote.paragraphs:
                                        fn_para = footnote.paragraphs[0]
                                        fn_para.style = doc.styles['Footnote Text']
                                        for run in fn_para.runs:
                                            run.font.name = 'Times New Roman'
                                            run.font.size = Pt(10)
                        else:
                            # Regular text
                            if part:
                                run = p.add_run(part)
                                if apply_formatting:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(12)
                else:
                    # No citation conversion, just add text as is
                    run = p.add_run(para_text)
                    if apply_formatting:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

            # Add page numbers in footer "1 of xx" only if formatting is requested
            if apply_formatting:
                section = doc.sections[0]
                footer = section.footer
                paragraph = footer.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = paragraph.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                add_page_number(run)
                
                run = paragraph.add_run(" of ")
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                
                run = paragraph.add_run()
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                add_num_pages(run)

            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            docx_content = buffer.getvalue()

            # Send response
            self.send_response(200)
            self.send_header('Content-type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            
            # Determine filename based on flags
            filename_parts = ["CitationFix"]
            if convert_citations:
                filename_parts.append("Converted")
            if apply_formatting:
                filename_parts.append("Formatted")
            
            filename = "-".join(filename_parts) + ".docx"
            
            self.send_header('Content-Disposition', f'attachment; filename="{filename}"')
            self.end_headers()
            self.wfile.write(docx_content)

        except Exception as e:
            self.send_response(500)
            self.send_header('Content-type', 'application/json')
            self.end_headers()
            error_msg = str(e)
            self.wfile.write(json.dumps({'error': 'Internal server error', 'details': error_msg}).encode())

    def do_GET(self):
        self.send_response(405)
        self.send_header('Content-type', 'application/json')
        self.end_headers()
        self.wfile.write(json.dumps({'error': 'Method not allowed'}).encode())
